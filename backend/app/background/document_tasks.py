import logging
import tempfile
import os
from typing import List

from langchain_community.document_loaders import PyPDFLoader
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document

from app.core.database import get_supabase_client
from app.features.knowledge.embedding import embed_texts

logger = logging.getLogger(__name__)

def extract_text_from_bytes(file_bytes: bytes, filename: str) -> List[Document]:
    """
    Extract text from PDF or DOCX bytes using Langchain loaders.
    Use temp files since loaders require file paths.
    """
    ext = filename.rsplit(".", 1)[-1].lower()
    docs = []
    
    with tempfile.NamedTemporaryFile(delete=False, suffix=f".{ext}") as temp_file:
        temp_file.write(file_bytes)
        temp_path = temp_file.name

    try:
        if ext == "pdf":
            loader = PyPDFLoader(temp_path)
            docs = loader.load()
        elif ext == "docx":
            # For DOCX, we can use docx manually or a loader.
            import docx
            doc = docx.Document(temp_path)
            full_text = []
            for para in doc.paragraphs:
                if para.text.strip():
                    full_text.append(para.text)
            
            # Since DOCX doesn't have inherent page numbers in python-docx easily,
            # we just create a single Document or one per paragraph.
            docs = [Document(page_content="\n".join(full_text), metadata={"page": 1})]
        elif ext == "txt":
            # Simple text read
            with open(temp_path, "r", encoding="utf-8") as f:
                content = f.read()
            docs = [Document(page_content=content, metadata={"page": 1})]
        else:
            raise ValueError(f"Unsupported file extension: {ext}")
            
    finally:
        if os.path.exists(temp_path):
            os.remove(temp_path)
            
    return docs

def process_document_pipeline(material_id: str, user_id: str, file_bytes: bytes, filename: str, content_type: str):
    """
    Background task to process an uploaded document:
    1. Extract text (with page numbers if PDF).
    2. Chunk text (1000 size, 200 overlap).
    3. Batch embed vectors.
    4. Save to `material_chunks` table.
    5. Update `study_materials` processing_status to 'success' or 'failed'.
    """
    db = get_supabase_client()
    logger.info(f"🚀 Starting background processing for material_id: {material_id} ({filename})")
    
    try:
        # 1. Trích xuất văn bản
        raw_docs = extract_text_from_bytes(file_bytes, filename)
        if not raw_docs:
            raise ValueError("No text could be extracted from the document.")
            
        # 2. Băm nhỏ văn bản (Chunking)
        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=1000,
            chunk_overlap=200,
            separators=["\n\n", "\n", " ", ""]
        )
        chunks = text_splitter.split_documents(raw_docs)
        logger.info(f"✅ Extracted {len(raw_docs)} pages, generated {len(chunks)} chunks.")
        
        # 3. Batch Embedding (với Tenacity @retry đã bọc trong hàm embed_texts)
        # Băm thành từng mẻ 50 chunks để tránh sập API
        BATCH_SIZE = 50
        all_vectors = []
        all_texts = [chunk.page_content for chunk in chunks]
        
        for i in range(0, len(all_texts), BATCH_SIZE):
            batch_texts = all_texts[i:i+BATCH_SIZE]
            logger.info(f"🔄 Embedding batch {i//BATCH_SIZE + 1}/{(len(all_texts)//BATCH_SIZE) + 1}...")
            batch_vectors = embed_texts(batch_texts)
            all_vectors.extend(batch_vectors)
            
        if len(all_vectors) != len(chunks):
            raise Exception("Mismatch between number of chunks and generated vectors.")
            
        # 4. Lưu vào Database (bảng material_chunks)
        insert_data = []
        for index, (chunk, vector) in enumerate(zip(chunks, all_vectors)):
            page_num = chunk.metadata.get("page", 0)  # PyPDFLoader returns 0-indexed page numbers usually
            
            insert_data.append({
                "material_id": material_id,
                "content": chunk.page_content,
                "chunk_index": index,
                "page_number": page_num + 1 if isinstance(page_num, int) else None, # Convert to 1-indexed
                "embedding": vector
            })
            
        # Lại phải insert theo batch để DB không bị ngợp thao tác mảng bự
        for i in range(0, len(insert_data), BATCH_SIZE):
            batch_insert = insert_data[i:i+BATCH_SIZE]
            db.table("material_chunks").insert(batch_insert).execute()
            
        logger.info(f"✅ Inserted {len(insert_data)} chunks into DB.")
        
        # 5. Cập nhật trạng thái thành công
        db.table("study_materials").update({
            "processing_status": "success"
        }).eq("id", material_id).execute()
        
        logger.info(f"🎉 Document pipeline finished successfully for {material_id}.")

    except Exception as e:
        logger.error(f"❌ Document pipeline failed for {material_id}: {str(e)}")
        # Cập nhật trạng thái lỗi để hiển thị lên Dashboard UI
        db.table("study_materials").update({
            "processing_status": "failed"
        }).eq("id", material_id).execute()


def delete_document_pipeline(material_id: str, file_url: str | None, user_id: str):
    """
    Background task to delete a document:
    1. Delete file from S3 (if file_url exists)
    2. Delete record from DB (which cascades to material_chunks)
    """
    db = get_supabase_client()
    logger.info(f"🗑️ Starting background deletion for material_id: {material_id}")
    
    try:
        # 1. Delete from S3
        if file_url and not file_url.startswith("http"):
             try:
                 db.storage.from_("knowledge-base").remove([file_url])
                 logger.info(f"✅ Removed file from S3: {file_url}")
             except Exception as e:
                 logger.warning(f"⚠️ Could not remove file {file_url} from S3: {e}")
                 # Continue to DB deletion even if S3 fails (e.g. file already deleted or ghost file)
        
        # 2. Delete from DB (CASCADE will delete material_chunks)
        db.table("study_materials").delete().eq("id", material_id).eq("user_id", user_id).execute()
        logger.info(f"✅ Successfully deleted material {material_id} from DB.")

    except Exception as e:
        logger.error(f"❌ Failed to delete material {material_id}: {str(e)}")

